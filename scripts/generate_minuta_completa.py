#!/usr/bin/env python3
"""Gera minutas de QA com a cobertura maxima possivel do formulario.

Produz dois .docx (ambos gitignored — padrao `Proposta_*.docx`) e valida a
integridade de cada um (zip valido, sem tags Jinja residuais, abre no
python-docx, paraIds unicos):

1. Proposta_Minuta_Escopo_Unico_Completa.docx
   - Contratante PJ, modalidade MISTA, escopo unico por modalidade.
   - TODAS as modalidades de honorarios marcadas (consultiva: 4; contenciosa: 4).
   - Exito ativo; horas extra no modo "senioridade"; SLA, Disposicoes, Despesas
     e taxa de manutencao ativos. Exercita o caminho de honorarios inline do
     template (onde vivem os campos justificados de texto livre).

2. Proposta_Minuta_Multi_Escopo_Completa.docx
   - Contratante PF, modalidade MISTA, 2 escopos consultivos + 2 contenciosos,
     forma de pagamento por escopo, CADA escopo com todas as modalidades e
     campos preenchidos. Horas extra no modo "horaFixa". Exercita o caminho de
     Subdocs (multi-escopo).

Uso: python scripts/generate_minuta_completa.py
"""
from __future__ import annotations

import re
import sys
import zipfile
from io import BytesIO
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document

from pmra.data_mapper import form_to_context
from pmra.schema import (
    AcaoRow,
    AtoProcessualRow,
    Contato,
    Contratante,
    DespesaItem,
    Despesas,
    Disposicoes,
    Endereco,
    Escopo,
    EscopoConsultivoItem,
    EscopoContenciosoItem,
    HonorariosConsultiva,
    HonorariosConsultivaModalidades,
    HonorariosContenciosa,
    HonorariosContenciosaModalidades,
    ProposalForm,
    SenioridadeRow,
)
from pmra.template_engine import render_proposal

_DOCS = (
    "word/document.xml", "word/header1.xml", "word/header2.xml", "word/header3.xml",
    "word/footer1.xml", "word/footer2.xml", "word/footer3.xml",
)
_JINJA_RESIDUO = ("{{", "}}", "{%", "%}")


def _senioridade(extra: bool = True) -> list[SenioridadeRow]:
    base = [
        SenioridadeRow(categoria="Sócio",                valor="R$ 1.200,00"),
        SenioridadeRow(categoria="Associado Sênior",     valor="R$ 950,00"),
        SenioridadeRow(categoria="Associado Pleno",      valor="R$ 750,00"),
        SenioridadeRow(categoria="Associado Júnior",     valor="R$ 500,00"),
        SenioridadeRow(categoria="Estagiário/Paralegal", valor="R$ 280,00"),
    ]
    if extra:
        base.append(SenioridadeRow(categoria="Consultor Especialista", valor="R$ 1.500,00"))
    return base


def _acoes() -> list[AcaoRow]:
    return [
        AcaoRow(natureza="Trabalhista", fase="Conhecimento", valor="R$ 5.500,00"),
        AcaoRow(natureza="Trabalhista", fase="Execução",     valor="R$ 3.800,00"),
        AcaoRow(natureza="Cível",       fase="Conhecimento", valor="R$ 7.200,00"),
        AcaoRow(natureza="Tributária",  fase="Conhecimento e recursal", valor="R$ 9.500,00"),
        AcaoRow(natureza="Consumidor",  fase="Conhecimento", valor="R$ 4.100,00"),
    ]


def _atos() -> list[AtoProcessualRow]:
    base = [
        AtoProcessualRow(ato="Petição Inicial",   descricao="Elaboração e protocolo da peça inaugural.",                            valor="R$ 2.500,00"),
        AtoProcessualRow(ato="Contestação",        descricao="Elaboração e protocolo de defesa em processo judicial.",              valor="R$ 2.200,00"),
        AtoProcessualRow(ato="Recurso Ordinário",  descricao="Elaboração e protocolo de recurso ordinário na Justiça do Trabalho.", valor="R$ 3.000,00"),
        AtoProcessualRow(ato="Agravo",             descricao="Elaboração e protocolo de agravo de instrumento ou agravo interno.",  valor="R$ 1.800,00"),
        AtoProcessualRow(ato="Apelação",           descricao="Elaboração e protocolo de apelação cível em 1º grau de jurisdição.",  valor="R$ 3.500,00"),
        AtoProcessualRow(ato="Recurso Especial / Extraordinário", descricao="Recursos excepcionais ao STJ e STF, inclusive agravos correlatos.", valor="R$ 6.500,00"),
        AtoProcessualRow(ato="Audiência",          descricao="Comparecimento, acompanhamento e atuação em audiência.",              valor="R$ 1.500,00"),
        AtoProcessualRow(ato="Diligência Externa", descricao="Diligência presencial em cartório, órgão público ou unidade do Contratante.", valor="R$ 850,00"),
        AtoProcessualRow(ato="Sustentação Oral",   descricao="Sustentação oral em sessão de julgamento de tribunal.",               valor="R$ 4.000,00"),
    ]
    return base


def _hon_cons_full() -> HonorariosConsultiva:
    """Honorarios consultivos com TODAS as 4 modalidades e todos os campos."""
    return HonorariosConsultiva(
        modalidades=HonorariosConsultivaModalidades(
            hora_senioridade=True, hora_fixa=True, fixo_mensal=True, valor_projeto=True,
        ),
        tabela_senioridade=_senioridade(),
        hora_fixa_valor="R$ 780,00",
        fixo_mensal_valor="R$ 18.500,00",
        fixo_mensal_cap="35 horas",
        fixo_mensal_excedente="R$ 650,00",
        valor_projeto_total="R$ 120.000,00",
        valor_projeto_cap_ativo=True,
        valor_projeto_cap="300 horas",
        valor_projeto_forma_pagamento=(
            "Entrada de 30% na assinatura e o saldo em 6 (seis) parcelas mensais "
            "iguais e sucessivas, vencendo a primeira 30 dias após o início dos trabalhos."
        ),
    )


def _hon_cont_full(horas_extra_modo: str = "senioridade") -> HonorariosContenciosa:
    """Honorarios contenciosos com TODAS as 4 modalidades e todos os campos."""
    return HonorariosContenciosa(
        modalidades=HonorariosContenciosaModalidades(
            valor_acao=True, valor_ato_processual=True, preco_mensal_massa=True, valor_projeto=True,
        ),
        tabela_acoes=_acoes(),
        tabela_atos=_atos(),
        preco_mensal_valor="R$ 12.000,00",
        preco_mensal_maximo_acoes="25",
        preco_mensal_maximo_acoes_extenso="vinte e cinco",
        preco_mensal_criterio_excedentes=(
            "Ações que excederem o número máximo coberto serão cobradas individualmente, "
            "conforme a tabela de Valor Mensal Por Processo acima, mediante prévia ciência ao Contratante."
        ),
        valor_projeto_total="R$ 95.000,00",
        valor_projeto_fases_cobertas=(
            "Fase de conhecimento e execução, em primeira e segunda instâncias (TRT/TRF), "
            "incluindo embargos de declaração, agravos e demais incidentes processuais correlatos."
        ),
        valor_projeto_forma_pagamento=(
            "40% no ato da assinatura; 30% após a sentença de primeiro grau; "
            "30% após o trânsito em julgado da decisão."
        ),
        exito_ativo=True,
        exito_percentual="15",
        horas_extra_escopo_modo=horas_extra_modo,
        horas_extra_senioridade=_senioridade(extra=False) if horas_extra_modo == "senioridade" else [],
        horas_extra_valor="R$ 600,00" if horas_extra_modo == "horaFixa" else "",
    )


_SLA = (
    "Demandas de baixa complexidade: até 3 dias úteis;\n"
    "Demandas de média complexidade: até 5 dias úteis;\n"
    "Demandas de alta complexidade (ou em língua estrangeira): entre 5 e 10 dias úteis.\n"
    "Projetos de altíssima complexidade ou alto volume: prazo a definir em conjunto com o Contratante.\n"
    "Solicitações de urgência fora do SLA acima poderão ser atendidas mediante alinhamento entre as Partes."
)

_DISPOSICOES = (
    "1. Cláusula de não solicitação: as Partes comprometem-se a não contratar colaboradores "
    "da outra Parte durante a vigência e por 12 (doze) meses após o término deste contrato.\n\n"
    "2. Confidencialidade reforçada: aplicar-se-á NDA específico para informações estratégicas, "
    "com prazo de confidencialidade de 5 (cinco) anos.\n\n"
    "3. Reajuste anual: os honorários serão reajustados anualmente pelo IPCA-IBGE acumulado, "
    "aplicado na data-base da assinatura."
)

_DESPESAS = [
    DespesaItem(categoria="Despesas Logísticas", descricao="Deslocamento em veículo próprio: R$ 2,50 por km e estacionamento, táxi, transporte por aplicativo, hospedagem e alimentação."),
    DespesaItem(categoria="Despesas Gerais",     descricao="Custas e emolumentos judiciais e administrativos, depósitos recursais, diligências externas simples (R$ 150,00 por diligência) e serviços de correspondentes."),
    DespesaItem(categoria="Despesas Periciais",  descricao="Honorários periciais e custos de assistente técnico, suportados conforme determinação judicial e prévia ciência ao Contratante."),
]


def build_form_escopo_unico() -> ProposalForm:
    """PJ, MISTA, escopo único por modalidade, todas as modalidades marcadas."""
    return ProposalForm(
        contratante=Contratante(
            tipo_pessoa="juridica",
            razao_social="Indústria e Comércio Horizonte S.A.",
            cnpj="12.345.678/0001-90",
            endereco=Endereco(
                logradouro="Avenida Paulista", numero="1578", bairro="Bela Vista",
                cep="01310-200", cidade="São Paulo", uf="SP",
            ),
            contato_nome="Dra. Maria Helena Carvalho",
            contatos=[
                Contato(telefone="(11) 98765-4321", email="maria.carvalho@horizonte.com.br"),
                Contato(telefone="(11) 3456-7890",  email="juridico@horizonte.com.br"),
            ],
        ),
        escopo=Escopo(
            modalidade="mista",
            atuacao_consultiva=(
                "Assessoria jurídica consultiva permanente nas áreas societária, contratual, "
                "trabalhista preventiva e regulatória, abrangendo a elaboração e revisão de atos "
                "societários e contratos, a emissão de pareceres, due diligence, governança "
                "corporativa, compliance e adequação à LGPD, com atendimento contínuo às demandas "
                "do Contratante por meio de canal dedicado."
            ),
            atuacao_contenciosa=(
                "Patrocínio e acompanhamento de demandas judiciais e administrativas nas esferas "
                "trabalhista, cível, tributária e do consumidor, em primeira e segunda instâncias e "
                "tribunais superiores, incluindo a prática de todos os atos processuais, "
                "comparecimento a audiências e sustentações orais."
            ),
            sla_ativo=True,
            sla_descricao=_SLA,
        ),
        honorarios_consultiva=_hon_cons_full(),
        honorarios_contenciosa=_hon_cont_full(horas_extra_modo="senioridade"),
        despesas=Despesas(
            tabela_despesas=[d.model_copy() for d in _DESPESAS],
            taxa_manutencao_ativa=True,
            taxa_manutencao_processual="R$ 75,00 por processo/mês",
        ),
        disposicoes=Disposicoes(ativo=True, descricao=_DISPOSICOES),
    )


def build_form_multi_escopo() -> ProposalForm:
    """PF, MISTA, 2 escopos consultivos + 2 contenciosos, forma por escopo, tudo marcado."""
    return ProposalForm(
        contratante=Contratante(
            tipo_pessoa="fisica",
            nome="João Pedro de Almeida Souza",
            cpf="123.456.789-09",
            endereco=Endereco(
                logradouro="Rua das Acácias", numero="245", bairro="Savassi",
                cep="30140-120", cidade="Belo Horizonte", uf="MG",
            ),
            contato_nome="João Pedro de Almeida Souza",
            contatos=[
                Contato(telefone="(31) 99876-5432", email="joao.souza@example.com"),
                Contato(telefone="(31) 3222-1100",  email="contato@joaosouza.adv.br"),
            ],
        ),
        escopo=Escopo(
            modalidade="mista",
            sla_ativo=True,
            sla_descricao=_SLA,
            escopos_consultivos=[
                EscopoConsultivoItem(
                    letra="A",
                    descricao=(
                        "Escopo consultivo societário e de M&A: estruturação de operações, elaboração "
                        "de atas e alterações contratuais, acordos de sócios, due diligence e pareceres."
                    ),
                    honorarios=_hon_cons_full(),
                ),
                EscopoConsultivoItem(
                    letra="B",
                    descricao=(
                        "Escopo consultivo contratual e de privacidade: revisão e negociação de contratos "
                        "comerciais, termos de uso, políticas de privacidade e programa de adequação à LGPD."
                    ),
                    honorarios=_hon_cons_full(),
                ),
            ],
            escopos_contenciosos=[
                EscopoContenciosoItem(
                    letra="A",
                    descricao=(
                        "Escopo contencioso trabalhista de massa: defesa em reclamações trabalhistas em "
                        "todas as instâncias e regiões, incluindo audiências e recursos."
                    ),
                    honorarios=_hon_cont_full(),
                ),
                EscopoContenciosoItem(
                    letra="B",
                    descricao=(
                        "Escopo contencioso cível e tributário estratégico: ações de alta complexidade "
                        "perante juízos estaduais e federais e tribunais superiores."
                    ),
                    honorarios=_hon_cont_full(),
                ),
            ],
            forma_pagamento_por_escopo_consultiva=True,
            forma_pagamento_por_escopo_contenciosa=True,
        ),
        honorarios_consultiva=HonorariosConsultiva(),
        honorarios_contenciosa=_hon_cont_full(horas_extra_modo="horaFixa"),
        despesas=Despesas(
            tabela_despesas=[d.model_copy() for d in _DESPESAS],
            taxa_manutencao_ativa=True,
            taxa_manutencao_processual="R$ 90,00 por processo/mês",
        ),
        disposicoes=Disposicoes(ativo=True, descricao=_DISPOSICOES),
    )


def validate_docx(data: bytes, label: str) -> dict:
    """Checa integridade do .docx. Levanta AssertionError se algo estiver errado."""
    assert data[:4] == b"PK\x03\x04", f"{label}: não é um zip/docx válido"

    with zipfile.ZipFile(BytesIO(data)) as zf:
        names = set(zf.namelist())
        assert "word/document.xml" in names, f"{label}: faltando word/document.xml"

        paraids: list[str] = []
        for doc_name in _DOCS:
            if doc_name not in names:
                continue
            xml = zf.read(doc_name).decode("utf-8")
            for marker in _JINJA_RESIDUO:
                assert marker not in xml, f"{label}: tag Jinja não renderizada {marker!r} em {doc_name}"
            paraids += re.findall(r'w14:paraId="([0-9A-Fa-f]+)"', xml)

        dup = {pid for pid in paraids if paraids.count(pid) > 1}
        assert not dup, f"{label}: paraIds duplicados (corrompe no Word): {sorted(dup)[:5]}"

    # Abre no python-docx — falha se o XML estiver malformado/corrompido.
    doc = Document(BytesIO(data))
    n_par = len(doc.paragraphs)
    n_tbl = len(doc.tables)
    # Conta parágrafos justificados (w:jc=both) — efeito do fix de justificação.
    body_xml = zipfile.ZipFile(BytesIO(data)).read("word/document.xml").decode("utf-8")
    n_just = body_xml.count('<w:jc w:val="both"/>')

    return {
        "bytes": len(data), "paragrafos": n_par, "tabelas": n_tbl,
        "paraids": len(paraids), "justificados": n_just,
    }


def main() -> None:
    alvos = [
        ("Proposta_Minuta_Escopo_Unico_Completa.docx", build_form_escopo_unico(),
         "PJ · MISTA · escopo único · todas as modalidades"),
        ("Proposta_Minuta_Multi_Escopo_Completa.docx", build_form_multi_escopo(),
         "PF · MISTA · 2+2 escopos · forma por escopo · todas as modalidades"),
    ]
    print("Gerando minutas de QA…\n")
    ok = True
    for nome, form, desc in alvos:
        try:
            data = render_proposal(form_to_context(form))
            rel = validate_docx(data, nome)
            Path(nome).write_bytes(data)
            print(f"  ✔ {nome}")
            print(f"     {desc}")
            print(f"     {rel['bytes']:,} bytes · {rel['paragrafos']} parágrafos · "
                  f"{rel['tabelas']} tabelas · {rel['justificados']} parágrafos justificados · "
                  f"{rel['paraids']} paraIds únicos\n")
        except AssertionError as e:
            ok = False
            print(f"  ✘ {nome}: FALHA DE INTEGRIDADE → {e}\n")

    print("Integridade OK — minutas prontas." if ok else "ATENÇÃO: houve falha de integridade.")
    sys.exit(0 if ok else 1)


if __name__ == "__main__":
    main()
